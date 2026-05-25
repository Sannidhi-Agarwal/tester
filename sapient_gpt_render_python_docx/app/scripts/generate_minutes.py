#!/usr/bin/env python3
"""generate_minutes.py - Sapient Finserv Minutes-of-the-Meeting renderer."""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Mm, Pt, RGBColor
from lxml import etree

BRAND_GREEN = "123D33"
LIGHT_GREEN = "C6E0B4"  # band colour for subtitle/Net Gain/XIRR
DISCLAIMER_GREY = "7F7F7F"
HEAD_FONT = "Aquatico"
BODY_FONT = "Gotham Medium"

PAGE_W_TWIPS = 11910
PAGE_H_TWIPS = 16850
LEFT_MARGIN_TWIPS = 1300
RIGHT_MARGIN_TWIPS = 720
TOP_MARGIN_TWIPS = 2160
BOTTOM_MARGIN_TWIPS = 2880
HEADER_DIST_TWIPS = 360
FOOTER_DIST_TWIPS = 360

SCRIPT_DIR = Path(__file__).resolve().parent
LOGO_CANDIDATES = [
    SCRIPT_DIR / "../assets/owl_logo.png",
    SCRIPT_DIR / "assets/owl_logo.png",
    SCRIPT_DIR / "owl_logo.png",
]


def find_logo():
    for c in LOGO_CANDIDATES:
        if c.exists():
            return c.resolve()
    return None


DECOR_SHAPES_XML = r"""<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
     xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
     xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
     xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
     xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:rPr><w:noProof/></w:rPr>
  <w:drawing>
    <wp:anchor distT="0" distB="0" distL="0" distR="0"
               simplePos="0" relativeHeight="1" behindDoc="1" locked="1"
               layoutInCell="1" allowOverlap="1">
      <wp:simplePos x="0" y="0"/>
      <wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>
      <wp:positionV relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionV>
      <wp:extent cx="280000" cy="10696575"/>
      <wp:effectExtent l="0" t="0" r="0" b="0"/>
      <wp:wrapNone/>
      <wp:docPr id="2000" name="LeftSidebar"/>
      <wp:cNvGraphicFramePr/>
      <a:graphic>
        <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
          <wps:wsp>
            <wps:cNvSpPr/>
            <wps:spPr>
              <a:xfrm><a:off x="0" y="0"/><a:ext cx="280000" cy="10696575"/></a:xfrm>
              <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
              <a:solidFill><a:srgbClr val="123D33"/></a:solidFill>
              <a:ln><a:noFill/></a:ln>
            </wps:spPr>
            <wps:bodyPr/>
          </wps:wsp>
        </a:graphicData>
      </a:graphic>
    </wp:anchor>
  </w:drawing>
  <w:drawing>
    <wp:anchor distT="0" distB="0" distL="114300" distR="114300"
               simplePos="0" relativeHeight="2" behindDoc="1" locked="1"
               layoutInCell="1" allowOverlap="1">
      <wp:simplePos x="0" y="0"/>
      <wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>
      <wp:positionV relativeFrom="page"><wp:posOffset>900000</wp:posOffset></wp:positionV>
      <wp:extent cx="574675" cy="1135380"/>
      <wp:effectExtent l="0" t="0" r="0" b="0"/>
      <wp:wrapNone/>
      <wp:docPr id="2001" name="ChevronLeft"/>
      <wp:cNvGraphicFramePr/>
      <a:graphic>
        <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
          <wps:wsp>
            <wps:cNvSpPr/>
            <wps:spPr>
              <a:xfrm><a:off x="0" y="0"/><a:ext cx="574675" cy="1135380"/></a:xfrm>
              <a:custGeom>
                <a:avLst/><a:gdLst/><a:ahLst/><a:cxnLst/>
                <a:rect l="0" t="0" r="r" b="b"/>
                <a:pathLst>
                  <a:path w="100" h="100">
                    <a:moveTo><a:pt x="0" y="0"/></a:moveTo>
                    <a:lnTo><a:pt x="100" y="50"/></a:lnTo>
                    <a:lnTo><a:pt x="0" y="100"/></a:lnTo>
                    <a:close/>
                  </a:path>
                </a:pathLst>
              </a:custGeom>
              <a:solidFill><a:srgbClr val="123D33"/></a:solidFill>
              <a:ln><a:noFill/></a:ln>
            </wps:spPr>
            <wps:bodyPr/>
          </wps:wsp>
        </a:graphicData>
      </a:graphic>
    </wp:anchor>
  </w:drawing>
  <w:drawing>
    <wp:anchor distT="0" distB="0" distL="114300" distR="114300"
               simplePos="0" relativeHeight="3" behindDoc="1" locked="1"
               layoutInCell="1" allowOverlap="1">
      <wp:simplePos x="0" y="0"/>
      <wp:positionH relativeFrom="page"><wp:posOffset>6988175</wp:posOffset></wp:positionH>
      <wp:positionV relativeFrom="page"><wp:posOffset>900000</wp:posOffset></wp:positionV>
      <wp:extent cx="574675" cy="1135380"/>
      <wp:effectExtent l="0" t="0" r="0" b="0"/>
      <wp:wrapNone/>
      <wp:docPr id="2002" name="ChevronRight"/>
      <wp:cNvGraphicFramePr/>
      <a:graphic>
        <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
          <wps:wsp>
            <wps:cNvSpPr/>
            <wps:spPr>
              <a:xfrm><a:off x="0" y="0"/><a:ext cx="574675" cy="1135380"/></a:xfrm>
              <a:custGeom>
                <a:avLst/><a:gdLst/><a:ahLst/><a:cxnLst/>
                <a:rect l="0" t="0" r="r" b="b"/>
                <a:pathLst>
                  <a:path w="100" h="100">
                    <a:moveTo><a:pt x="100" y="0"/></a:moveTo>
                    <a:lnTo><a:pt x="0" y="50"/></a:lnTo>
                    <a:lnTo><a:pt x="100" y="100"/></a:lnTo>
                    <a:close/>
                  </a:path>
                </a:pathLst>
              </a:custGeom>
              <a:solidFill><a:srgbClr val="123D33"/></a:solidFill>
              <a:ln><a:noFill/></a:ln>
            </wps:spPr>
            <wps:bodyPr/>
          </wps:wsp>
        </a:graphicData>
      </a:graphic>
    </wp:anchor>
  </w:drawing>
  <w:drawing>
    <wp:anchor distT="0" distB="0" distL="114300" distR="114300"
               simplePos="0" relativeHeight="4" behindDoc="1" locked="1"
               layoutInCell="1" allowOverlap="1">
      <wp:simplePos x="0" y="0"/>
      <wp:positionH relativeFrom="page"><wp:posOffset>6700000</wp:posOffset></wp:positionH>
      <wp:positionV relativeFrom="page"><wp:posOffset>9300000</wp:posOffset></wp:positionV>
      <wp:extent cx="2247900" cy="2070100"/>
      <wp:effectExtent l="0" t="0" r="0" b="0"/>
      <wp:wrapNone/>
      <wp:docPr id="2003" name="CornerDonut"/>
      <wp:cNvGraphicFramePr/>
      <a:graphic>
        <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
          <wps:wsp>
            <wps:cNvSpPr/>
            <wps:spPr>
              <a:xfrm><a:off x="0" y="0"/><a:ext cx="2247900" cy="2070100"/></a:xfrm>
              <a:prstGeom prst="donut"><a:avLst><a:gd name="adj" fmla="val 18000"/></a:avLst></a:prstGeom>
              <a:solidFill><a:srgbClr val="123D33"/></a:solidFill>
              <a:ln><a:noFill/></a:ln>
            </wps:spPr>
            <wps:bodyPr/>
          </wps:wsp>
        </a:graphicData>
      </a:graphic>
    </wp:anchor>
  </w:drawing>
</w:r>
"""
def add_decorations_to_paragraph(paragraph):
    run_xml = etree.fromstring(DECOR_SHAPES_XML)
    paragraph._p.append(run_xml)


BOLD_RE = re.compile(r"\*\*(.+?)\*\*", re.DOTALL)


def add_inline_runs(paragraph, text, *, font=BODY_FONT, size_pt=9.5,
                    color=BRAND_GREEN, bold_default=False):
    cursor = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > cursor:
            _run(paragraph, text[cursor:m.start()], font, size_pt, color,
                 bold=bold_default)
        _run(paragraph, m.group(1), font, size_pt, color, bold=True)
        cursor = m.end()
    if cursor < len(text):
        _run(paragraph, text[cursor:], font, size_pt, color, bold=bold_default)


def _run(paragraph, text, font, size_pt, color,
         *, bold=False, italic=False, underline=False):
    if not text:
        return None
    run = paragraph.add_run(text)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    return run


def set_paragraph_spacing(paragraph, *, before_pt=0, after_pt=0,
                          line_pt=None, line_rule="auto"):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(int(before_pt * 20)))
    spacing.set(qn("w:after"), str(int(after_pt * 20)))
    if line_pt is not None:
        spacing.set(qn("w:line"), str(int(line_pt * 20)))
        spacing.set(qn("w:lineRule"), line_rule)


def set_paragraph_indent(paragraph, *, left_twips=0, hanging_twips=0,
                         first_twips=0):
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    if left_twips:
        ind.set(qn("w:left"), str(left_twips))
    if hanging_twips:
        ind.set(qn("w:hanging"), str(hanging_twips))
    if first_twips:
        ind.set(qn("w:firstLine"), str(first_twips))


def configure_section(section):
    section.page_width = Emu(int(PAGE_W_TWIPS * 635))
    section.page_height = Emu(int(PAGE_H_TWIPS * 635))
    section.left_margin = Emu(int(LEFT_MARGIN_TWIPS * 635))
    section.right_margin = Emu(int(RIGHT_MARGIN_TWIPS * 635))
    section.top_margin = Emu(int(TOP_MARGIN_TWIPS * 635))
    section.bottom_margin = Emu(int(BOTTOM_MARGIN_TWIPS * 635))
    section.header_distance = Emu(int(HEADER_DIST_TWIPS * 635))
    section.footer_distance = Emu(int(FOOTER_DIST_TWIPS * 635))


def build_header(section, logo_path):
    header = section.header
    header.is_linked_to_previous = False
    p_decor = header.paragraphs[0]
    add_decorations_to_paragraph(p_decor)
    if logo_path and logo_path.exists():
        p_logo = header.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p_logo, before_pt=0, after_pt=0)
        run = p_logo.add_run()
        run.add_picture(str(logo_path), width=Mm(20))
    header.add_paragraph()


DISCLAIMER_TEXT = (
    "Disclaimer: This document is meant for private circulation only. "
    "Sapient Finserv Pvt. Ltd has taken due care while compiling this report. "
    "All information/opinion contained/expressed herein above by Sapient Finserv "
    "has been based upon information available to the sources, we believe to be "
    "reliable, but we do not make any representation or warranty as to its "
    "accuracy, completeness or correctness. Readers should use this information "
    "at their own risk. Sapient Finserv shall not be held responsible for any "
    "direct or indirect loss caused by relying on this information. Mutual Funds "
    "and securities investments are subject to market risks and there is no "
    "assurance or guarantee that the objectives of the Scheme will be achieved. "
    "Past performance of the Sponsor/ Mutual Fund/Investment Manager are not "
    "indicative of the future performance of the Scheme(s)."
)


def build_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p_amfi = footer.paragraphs[0]
    p_amfi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p_amfi, before_pt=0, after_pt=2)
    _run(p_amfi, "AMFI Registered Mutual Fund Distributor",
         font=BODY_FONT, size_pt=9, color=BRAND_GREEN, bold=True)
    p_disc = footer.add_paragraph()
    p_disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p_disc, before_pt=0, after_pt=0,
                          line_pt=10, line_rule="auto")
    _run(p_disc, DISCLAIMER_TEXT, font=BODY_FONT, size_pt=6,
         color=DISCLAIMER_GREY)


def add_title(doc, title="Minutes of the Meeting"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before_pt=0, after_pt=8)
    _run(p, title, font=HEAD_FONT, size_pt=15, color=BRAND_GREEN, bold=False)


def add_meeting_info(doc, data):
    info_lines = [
        ("Client", data.get("client", "")),
        ("Date", data.get("date", "")),
        ("Attendees", data.get("attendees", "")),
        ("Location", data.get("location", "")),
    ]
    for label, value in info_lines:
        if not value:
            continue
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before_pt=0, after_pt=2, line_pt=12)
        _run(p, label + ": ", font=BODY_FONT, size_pt=10,
             color=BRAND_GREEN, bold=True)
        add_inline_runs(p, str(value), font=BODY_FONT, size_pt=10,
                        color=BRAND_GREEN)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before_pt=0, after_pt=4)


def add_section_title(doc, title):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=8, after_pt=4)
    _run(p, title, font=BODY_FONT, size_pt=11, color=BRAND_GREEN,
         bold=True, underline=True)


def add_bullet(doc, text, *, indent_twips=360):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=0, after_pt=2, line_pt=12)
    set_paragraph_indent(p, left_twips=indent_twips, hanging_twips=200)
    _run(p, "•  ", font=BODY_FONT, size_pt=9.5, color=BRAND_GREEN)
    add_inline_runs(p, text, font=BODY_FONT, size_pt=9.5, color=BRAND_GREEN)


def add_paragraph_text(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=2, after_pt=2, line_pt=13)
    add_inline_runs(p, text, font=BODY_FONT, size_pt=10, color=BRAND_GREEN)


def render_sip(doc, sec):
    title = sec.get("title", "Systematic Investment Plan (SIP)")
    add_section_title(doc, title)
    for line in sec.get("preamble", []):
        add_bullet(doc, line)
    intro = sec.get("intro")
    if intro:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before_pt=4, after_pt=2)
        add_inline_runs(p, intro, font=BODY_FONT, size_pt=10,
                        color=BRAND_GREEN, bold_default=True)
    columns = sec.get("columns") or [[]]
    if len(columns) == 1:
        columns = [columns[0], []]
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement("w:" + edge)
        b.set(qn("w:val"), "nil")
        borders.append(b)
    tblPr.append(borders)
    avail = PAGE_W_TWIPS - LEFT_MARGIN_TWIPS - RIGHT_MARGIN_TWIPS
    col_w = avail // 2
    for cell, lines in zip(table.rows[0].cells, columns):
        cell.width = Emu(int(col_w * 635))
        first_p = cell.paragraphs[0]
        first_p._element.getparent().remove(first_p._element)
        for line in lines:
            p = cell.add_paragraph()
            set_paragraph_spacing(p, before_pt=0, after_pt=2, line_pt=12)
            stripped = line.lstrip()
            if stripped.startswith("o ") and len(line) - len(stripped) >= 2:
                set_paragraph_indent(p, left_twips=360, hanging_twips=180)
                _run(p, "o  ", font=BODY_FONT, size_pt=9.5, color=BRAND_GREEN)
                add_inline_runs(p, stripped[2:], font=BODY_FONT, size_pt=9.5,
                                color=BRAND_GREEN)
            else:
                add_inline_runs(p, line, font=BODY_FONT, size_pt=9.5,
                                color=BRAND_GREEN)


def render_groups(doc, sec):
    add_section_title(doc, sec.get("title", ""))
    for group in sec.get("groups", []):
        gtitle = group.get("title")
        if gtitle:
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before_pt=4, after_pt=2)
            _run(p, gtitle, font=BODY_FONT, size_pt=10,
                 color=BRAND_GREEN, bold=True)
        for block in group.get("blocks", []):
            heading = block.get("heading")
            if heading:
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before_pt=2, after_pt=2)
                set_paragraph_indent(p, left_twips=180)
                _run(p, heading, font=BODY_FONT, size_pt=10,
                     color=BRAND_GREEN, bold=True)
            for bullet in block.get("bullets", []):
                add_bullet(doc, bullet, indent_twips=540)


def render_section(doc, sec):
    stype = sec.get("type", "bullets")
    if stype == "bullets":
        if sec.get("title"):
            add_section_title(doc, sec["title"])
        for item in sec.get("items", []):
            add_bullet(doc, item)
    elif stype == "paragraph":
        if sec.get("title"):
            add_section_title(doc, sec["title"])
        for para in sec.get("paragraphs", []) or sec.get("items", []):
            add_paragraph_text(doc, para)
    elif stype == "groups":
        render_groups(doc, sec)
    elif stype == "sip":
        render_sip(doc, sec)
    else:
        if sec.get("title"):
            add_section_title(doc, sec["title"])
        for item in sec.get("items", []):
            add_bullet(doc, item)


PORTFOLIO_DEFAULT = {
    "title":    "Journey So Far",
    "subtitle": "Reporting of Portfolio Performance/ Other Investments",
    "intro":    ("In our discussion, we had reported the performance of "
                 "your investments with us. A synopsis of the same has "
                 "been mentioned below:"),
    "header":   ["Since Inception", "Equity", "Hybrid", "Debt & Others", "Total"],
    "rows": [
        ["Opening Balance :", "", "", "", ""],
        ["Purchase",          "", "", "", ""],
        ["Switch In",         "", "", "", ""],
        ["Switch Out",        "", "", "", ""],
        ["Div. PayOut",       "", "", "", ""],
        ["Redemption",        "", "", "", ""],
        ["Net Addition",      "", "", "", ""],
        ["Closing Balance",   "", "", "", ""],
        ["Net Gain :",        "", "", "", ""],
        ["XIRR(%) :",         "", "", "", ""],
    ],
}

LIGHT_BAND_ROWS = {"Net Gain :", "XIRR(%) :"}
BOLD_LABEL_ROWS = {"Opening Balance :", "Net Gain :", "XIRR(%) :"}


def _shade_paragraph(paragraph, color_hex):
    """Apply a solid background fill to a paragraph (used for banded rows)."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)


def _set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _set_cell_borders(cell, color_hex="BFBFBF", size="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement("w:" + edge)
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color_hex)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _set_cell_vertical_align(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    tcPr.append(vAlign)


def _set_page_break_before(paragraph):
    """Set w:pageBreakBefore on the paragraph so it always starts a new page.

    This is cleaner than inserting an empty paragraph with a page-break run,
    which can leave a blank page behind on some renderers.
    """
    pPr = paragraph._p.get_or_add_pPr()
    pbb = pPr.find(qn("w:pageBreakBefore"))
    if pbb is None:
        pbb = OxmlElement("w:pageBreakBefore")
        pPr.insert(0, pbb)


def render_portfolio_table(doc, table_data):
    p_h = doc.add_paragraph()
    p_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p_h, before_pt=0, after_pt=4)
    _set_page_break_before(p_h)
    _run(p_h, table_data.get("title", "Journey So Far"),
         font=HEAD_FONT, size_pt=15, color=BRAND_GREEN, bold=True)
    subtitle = table_data.get("subtitle")
    if subtitle:
        p_s = doc.add_paragraph()
        p_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p_s, before_pt=0, after_pt=4)
        _shade_paragraph(p_s, LIGHT_GREEN)
        _run(p_s, subtitle, font=BODY_FONT, size_pt=10,
             color=BRAND_GREEN, bold=True)
    intro = table_data.get("intro")
    if intro:
        p_i = doc.add_paragraph()
        set_paragraph_spacing(p_i, before_pt=0, after_pt=8, line_pt=14)
        add_inline_runs(p_i, intro, font=BODY_FONT, size_pt=10,
                        color=BRAND_GREEN)
    header = table_data.get("header") or PORTFOLIO_DEFAULT["header"]
    rows = table_data.get("rows") or PORTFOLIO_DEFAULT["rows"]
    n_cols = len(header)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.autofit = False
    avail = PAGE_W_TWIPS - LEFT_MARGIN_TWIPS - RIGHT_MARGIN_TWIPS
    first_w = int(avail * 0.26)
    other_w = (avail - first_w) // (n_cols - 1)
    widths = [first_w] + [other_w] * (n_cols - 1)
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    hdr_row = table.rows[0]
    for i, cell in enumerate(hdr_row.cells):
        cell.width = Emu(int(widths[i] * 635))
        _set_cell_shading(cell, BRAND_GREEN)
        _set_cell_borders(cell, color_hex=BRAND_GREEN, size="4")
        _set_cell_vertical_align(cell)
        first_p = cell.paragraphs[0]
        first_p.text = ""
        first_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_spacing(first_p, before_pt=4, after_pt=4)
        _run(first_p, header[i], font=BODY_FONT, size_pt=10,
             color="FFFFFF", bold=True)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        label = row_data[0] if row_data else ""
        is_band = label in LIGHT_BAND_ROWS
        is_bold_label = label in BOLD_LABEL_ROWS
        for c_idx, cell in enumerate(row.cells):
            cell.width = Emu(int(widths[c_idx] * 635))
            _set_cell_borders(cell, color_hex="BFBFBF", size="4")
            _set_cell_vertical_align(cell)
            if is_band:
                _set_cell_shading(cell, LIGHT_GREEN)
            text_color = BRAND_GREEN
            bold = is_bold_label
            value = row_data[c_idx] if c_idx < len(row_data) else ""
            first_p = cell.paragraphs[0]
            first_p.text = ""
            # Right-align EVERY cell to match the original Format.docx layout
            first_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_paragraph_spacing(first_p, before_pt=3, after_pt=3, line_pt=13)
            _run(first_p, value, font=BODY_FONT, size_pt=10,
                 color=text_color, bold=bold)


def build_document(data, output_path, logo_path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor.from_string(BRAND_GREEN)
    section = doc.sections[0]
    configure_section(section)
    build_header(section, logo_path)
    build_footer(section)
    add_title(doc)
    add_meeting_info(doc, data)
    for sec in data.get("sections", []):
        render_section(doc, sec)
    table_data = dict(PORTFOLIO_DEFAULT)
    overrides = data.get("portfolio_table") or {}
    table_data.update({k: v for k, v in overrides.items() if v is not None})
    render_portfolio_table(doc, table_data)
    doc.save(str(output_path))


def main(argv):
    if len(argv) != 3:
        print("Usage: generate_minutes.py <input.json> <output.docx>",
              file=sys.stderr)
        return 1
    in_path = Path(argv[1])
    out_path = Path(argv[2])
    if not in_path.exists():
        print("Input not found: " + str(in_path), file=sys.stderr)
        return 2
    data = json.loads(in_path.read_text(encoding="utf-8"))
    logo = find_logo()
    if logo is None:
        print("Warning: owl_logo.png not found - header will have no logo.",
              file=sys.stderr)
    build_document(data, out_path, logo)
    print("Wrote " + str(out_path))
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))
